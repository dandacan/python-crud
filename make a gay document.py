from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
RUSure = raw_input("do you really want to waste loads of time making the stupid file? (yes or no)")
if RUSure == "yes" or RUSure == "Yes":
    doc = Document()
    for r in xrange(0,64):
    	paragraph = doc.add_paragraph()
        for g in xrange(0,64):
            for b in xrange(0,64):
            	if (b % 2 == 0):
            	    text = "ur mom gay"
            	else:
            		text = "no u"
                run = paragraph.add_run(" \n " + text)
                font = run.font
                font.color.rgb = RGBColor(r * 4, g * 4, b * 4)
            print "done with " + str(r) + " " + str(g)
        doc.save('gay.docx')